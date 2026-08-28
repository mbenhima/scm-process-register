# -*- coding: utf-8 -*-
import warnings
import docx
from docx.shared import RGBColor, Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
HEADER_FILL = '1F4B45'

d = docx.Document('fw-pandoc.docx')

# ---- Move the native TOC field from the very top of the document (where pandoc
# ---- puts it) to after the title block, so the reading order is Cover -> TOC ->
# ---- Part 0, not TOC -> Cover. Page breaks pin the cover to its own page and the
# ---- TOC to its own page range regardless of how long either renders.
def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p

body = d.element.body
sdt = body.find(qn('w:sdt'))
if sdt is not None:
    body.remove(sdt)
    title_end_p = d.paragraphs[4]._p  # 'Version ...' -- last title-block paragraph
    pb_before_toc = make_page_break_para()
    pb_after_toc = make_page_break_para()
    title_end_p.addnext(pb_before_toc)
    pb_before_toc.addnext(sdt)
    sdt.addnext(pb_after_toc)

# ---- Margins: more usable width for the wider reference tables ----
for section in d.sections:
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.page_height is None:
        section.page_height = Inches(11)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
USABLE_WIDTH = d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin

# ---- Explicit column widths per table type ----
OVERVIEW_WIDTHS = [0.16, 0.30, 0.30, 0.24]          # Framework, Altitude, Question it answers, Logged on
INTERACTION_MAP_WIDTHS = [0.10, 0.17, 0.15, 0.17, 0.41]  # Lewin Stage, ADKAR Focus, Bridges, Kübler-Ross, What's Happening
SIGNALS_WIDTHS = [0.12, 0.22, 0.12, 0.54]           # ID, Name, Type, Detailed Description
DECISION_MATRIX_WIDTHS = [0.14, 0.13, 0.24, 0.23, 0.26]  # Framework, Transition, Must-Have, Nice-to-Have, Who Decides
KPI_WIDTHS = [0.16, 0.84]                           # Field, Value
APPENDIX_WIDTHS = [0.24, 0.76]                      # ID Range, What It Covers


def header_texts(t):
    return [c.text.strip() for c in t.rows[0].cells]


def is_overview_table(t):
    h = header_texts(t)
    return len(t.columns) == 4 and h[0] == 'Framework' and h[1] == 'Altitude'


def is_interaction_map_table(t):
    h = header_texts(t)
    return len(t.columns) == 5 and h[0] == 'Lewin Stage'


def is_signals_table(t):
    h = header_texts(t)
    return len(t.columns) == 4 and h[0] == 'ID' and h[1] == 'Name' and h[2] == 'Type'


def is_decision_matrix_table(t):
    h = header_texts(t)
    return len(t.columns) == 5 and h[0] == 'Framework' and h[1] == 'Transition'


def is_kpi_table(t):
    h = header_texts(t)
    return len(t.columns) == 2 and h[0] == 'Field' and h[1] == 'Value'


def is_appendix_table(t):
    h = header_texts(t)
    return len(t.columns) == 2 and h[0] == 'ID Range'


def set_col_widths(table, ratios, total=None):
    total = total if total is not None else USABLE_WIDTH
    widths = [Emu(int(total * r)) for r in ratios]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol, w in zip(tblGrid.findall(qn('w:gridCol')), widths):
            gridCol.set(qn('w:w'), str(w))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(w / 635)))  # EMU -> twentieths of a point (dxa)


table_counts = {'overview': 0, 'interaction': 0, 'signals': 0, 'matrix': 0, 'kpi': 0, 'appendix': 0, 'unmatched': 0}
for t in d.tables:
    if is_overview_table(t):
        set_col_widths(t, OVERVIEW_WIDTHS)
        table_counts['overview'] += 1
    elif is_interaction_map_table(t):
        set_col_widths(t, INTERACTION_MAP_WIDTHS)
        table_counts['interaction'] += 1
    elif is_signals_table(t):
        set_col_widths(t, SIGNALS_WIDTHS)
        table_counts['signals'] += 1
    elif is_decision_matrix_table(t):
        set_col_widths(t, DECISION_MATRIX_WIDTHS)
        table_counts['matrix'] += 1
    elif is_kpi_table(t):
        set_col_widths(t, KPI_WIDTHS)
        table_counts['kpi'] += 1
    elif is_appendix_table(t):
        set_col_widths(t, APPENDIX_WIDTHS)
        table_counts['appendix'] += 1
    else:
        table_counts['unmatched'] += 1
print('table detector matches:', table_counts)

# ---- Prevent a table row's content from splitting across a page break, except
# ---- the KPI Field|Value tables, whose "Why this design" / "Worked example"
# ---- rows are long prose and would otherwise force near-one-row-per-page gaps.
for t in d.tables:
    if is_kpi_table(t):
        continue
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

# ---- Title block (paragraphs 0-4) ----
titles = d.paragraphs[0:5]
sizes = [30, 15, 11.5, 11.5, 11]
bolds = [True, False, False, False, False]
italics = [False, False, True, True, False]
for p, size, bold, italic in zip(titles, sizes, bolds, italics):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = H1_COLOR if size == 30 else RGBColor(0x33, 0x33, 0x33)
d.paragraphs[0].space_before = Pt(100)
d.paragraphs[4].space_after = Pt(20)

# ---- Heading colors and sizes ----
H4_COLOR = RGBColor(0x5A, 0x8F, 0x89)
HEADING_SPECS = {
    'Heading 1': (H1_COLOR, Pt(21)),
    'Heading 2': (H2_COLOR, Pt(17)),
    'Heading 3': (H3_COLOR, Pt(14.5)),
    'Heading 4': (H4_COLOR, Pt(13)),
}
for p in d.paragraphs:
    spec = HEADING_SPECS.get(p.style.name)
    if spec:
        color, size = spec
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color
            r.font.size = size
        if p.style.name == 'Heading 3':
            p.space_before = Pt(14)

# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Shade the "Field" column (col 0) of each KPI table for scan-ability,
# ---- since these tables run vertically (one KPI per table) rather than in rows ----
KPI_FIELD_FILL = 'DCEAE7'
for t in d.tables:
    if not is_kpi_table(t):
        continue
    for row in t.rows[1:]:
        shade_cell(row.cells[0], KPI_FIELD_FILL)

# ---- Base body font size increase (readability) ----
BODY_SIZE = Pt(13)
TABLE_SIZE = Pt(11.5)
HEADING_STYLE_NAMES = {'Heading 1', 'Heading 2', 'Heading 3', 'Title'}

for style_name in ('Normal', 'Body Text', 'Compact', 'First Paragraph', 'List Paragraph', 'Block Text'):
    if style_name in [s.name for s in d.styles]:
        d.styles[style_name].font.size = BODY_SIZE

for i, p in enumerate(d.paragraphs):
    if i < 5:
        continue  # title block already sized explicitly
    if p.style.name in HEADING_STYLE_NAMES:
        continue  # headings sized explicitly above
    for r in p.runs:
        if r.font.size is None:
            r.font.size = BODY_SIZE

for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = TABLE_SIZE

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — Change Management Frameworks'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

# ---- Force Word/LibreOffice to recompute TOC / PAGE fields automatically ----
settings = d.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

d.save('fw-styled.docx')
print('saved fw-styled.docx')
