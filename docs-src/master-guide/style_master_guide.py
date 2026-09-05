# -*- coding: utf-8 -*-
import copy
import warnings
import docx
from docx.shared import RGBColor, Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
HEADER_FILL = '1F4B45'

d = docx.Document('mg-pandoc.docx')

# ---- Move the native TOC field from the very top of the document (where pandoc
# ---- puts it) to after the title block, so the reading order is Cover -> TOC ->
# ---- Part 0, not TOC -> Cover. Page breaks pin the cover to its own page and the
# ---- TOC to its own page range regardless of how long either renders.
def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p

body = d.element.body
sdt = body.find(qn('w:sdt'))
if sdt is not None:
    body.remove(sdt)
    title_end_p = d.paragraphs[4]._p  # 'Version 1.0 ...' -- last title-block paragraph
    pb_before_toc = make_page_break_para()
    pb_after_toc = make_page_break_para()
    title_end_p.addnext(pb_before_toc)
    pb_before_toc.addnext(sdt)
    sdt.addnext(pb_after_toc)

# ---- Narrower margins: more usable width for wide reference tables ----
for section in d.sections:
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.page_height is None:
        section.page_height = Inches(11)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
USABLE_WIDTH = d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin
# Part 1B.2-1B.3 renders in a landscape section (see the section-split below) —
# same margins, swapped page dimensions, so its tables get ~34% more width.
LANDSCAPE_USABLE_WIDTH = d.sections[0].page_height - d.sections[0].left_margin - d.sections[0].right_margin

# ---- Explicit column widths for the 6-column alert-reference tables ----
ALERT_TABLE_WIDTHS = [0.09, 0.15, 0.08, 0.38, 0.15, 0.15]  # ID, Name, Severity, Trigger condition, Escalation, SLA
ALERT_NONLIVE_WIDTHS = [0.09, 0.28, 0.63]  # ID, Name, Why it never fires here
WBS_GANTT_WIDTHS = [0.11, 0.32, 0.06, 0.04, 0.07, 0.08, 0.13, 0.11, 0.08]  # ID, Task/Step Name, Track, Ph., Week(s), Lewin, ADKAR, Bridges, Kubler-Ross
CALENDAR_WIDTHS = [0.07, 0.16, 0.12, 0.16, 0.14, 0.16, 0.19]  # Week, Phase(s) Active, Lewin, ADKAR Focus, Bridges, Kubler-Ross, Exception
SPRINT_INDEX_WIDTHS = [0.08, 0.08, 0.38, 0.24, 0.22]  # Sprint, Weeks, Scope of Work, Justifying Principle(s), Task/Step ID
BUILD_INDEX_WIDTHS = [0.10, 0.14, 0.08, 0.52, 0.16]  # ID, Phase, Week(s), Task, Owner
CHECKLIST_WIDTHS = [0.20, 0.10, 0.55, 0.15]  # Phase, Track, Checklist Item, Weight %
CHARTER_MATRIX_WIDTHS = [0.28, 0.08, 0.12, 0.30, 0.22]  # Charter, Phase, Week(s), Task/Step ID, CRUD Action


def is_wbs_gantt_table(t):
    return len(t.columns) == 9 and t.rows[0].cells[0].text.strip() == 'ID'


def is_calendar_table(t):
    return len(t.columns) == 7 and t.rows[0].cells[0].text.strip() == 'Week'


def is_sprint_index_table(t):
    return len(t.columns) == 5 and t.rows[0].cells[0].text.strip() == 'Sprint'


def is_build_index_table(t):
    return len(t.columns) == 5 and t.rows[0].cells[0].text.strip() == 'ID' and t.rows[0].cells[1].text.strip() == 'Phase'


def is_checklist_table(t):
    return len(t.columns) == 4 and t.rows[0].cells[0].text.strip() == 'Phase' and t.rows[0].cells[3].text.strip() == 'Weight %'


def is_charter_matrix_table(t):
    return len(t.columns) == 5 and t.rows[0].cells[0].text.strip() == 'Charter' and t.rows[0].cells[4].text.strip() == 'CRUD Action'


def set_col_widths(table, ratios, total=None):
    total = total if total is not None else USABLE_WIDTH
    widths = [Emu(int(total * r)) for r in ratios]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol, w in zip(tblGrid.findall(qn('w:gridCol')), widths):
            gridCol.set(qn('w:w'), str(w))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(w / 635)))  # EMU -> twentieths of a point (dxa)


WEEKLY_TABLE_WIDTHS = [0.09, 0.24, 0.24, 0.31, 0.12]  # Week, PM Track, CM Track, journi Entry, Exception


def is_weekly_track_table(t):
    return len(t.columns) == 5 and 'PM) Track' in t.rows[0].cells[1].text

for t in d.tables:
    if len(t.columns) == 6 and t.rows[0].cells[0].text.strip() == 'ID':
        set_col_widths(t, ALERT_TABLE_WIDTHS)
    elif len(t.columns) == 3 and t.rows[0].cells[0].text.strip() == 'ID':
        set_col_widths(t, ALERT_NONLIVE_WIDTHS)
    elif is_wbs_gantt_table(t):
        set_col_widths(t, WBS_GANTT_WIDTHS, total=LANDSCAPE_USABLE_WIDTH)
    elif is_calendar_table(t):
        set_col_widths(t, CALENDAR_WIDTHS, total=LANDSCAPE_USABLE_WIDTH)
    elif is_sprint_index_table(t):
        set_col_widths(t, SPRINT_INDEX_WIDTHS, total=LANDSCAPE_USABLE_WIDTH)
    elif is_build_index_table(t):
        set_col_widths(t, BUILD_INDEX_WIDTHS, total=LANDSCAPE_USABLE_WIDTH)
    elif is_weekly_track_table(t):
        set_col_widths(t, WEEKLY_TABLE_WIDTHS, total=LANDSCAPE_USABLE_WIDTH)
    elif is_checklist_table(t):
        set_col_widths(t, CHECKLIST_WIDTHS)
    elif is_charter_matrix_table(t):
        set_col_widths(t, CHARTER_MATRIX_WIDTHS)

# ---- Prevent a table row's content from splitting across a page break ----
# Skip the large PM/CM weekly-track tables: their cells are prose-length, and
# forcing cantSplit there produces near-one-row-per-page with heavy blank space.
for t in d.tables:
    if is_weekly_track_table(t):
        continue
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

# ---- Title block (paragraphs 0-4) ----
titles = d.paragraphs[0:5]
sizes = [30, 15, 11.5, 11.5, 11]
bolds = [True, False, False, False, False]
italics = [False, False, True, True, False]
for p, size, bold, italic in zip(titles, sizes, bolds, italics):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = H1_COLOR if size == 30 else RGBColor(0x33, 0x33, 0x33)
d.paragraphs[0].space_before = Pt(100)
d.paragraphs[4].space_after = Pt(20)

# ---- Heading colors and sizes (bumped up to stay ahead of the larger body text) ----
H4_COLOR = RGBColor(0x5A, 0x8F, 0x89)
HEADING_SPECS = {
    'Heading 1': (H1_COLOR, Pt(21)),
    'Heading 2': (H2_COLOR, Pt(17)),
    'Heading 3': (H3_COLOR, Pt(14.5)),
    'Heading 4': (H4_COLOR, Pt(13)),
}
for p in d.paragraphs:
    spec = HEADING_SPECS.get(p.style.name)
    if spec:
        color, size = spec
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color
            r.font.size = size
        if p.style.name == 'Heading 3':
            p.space_before = Pt(14)

# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Base body font size increase (readability) ----
BODY_SIZE = Pt(13)
TABLE_SIZE = Pt(11.5)
HEADING_STYLE_NAMES = {'Heading 1', 'Heading 2', 'Heading 3', 'Title'}

for style_name in ('Normal', 'Body Text', 'Compact', 'First Paragraph', 'List Paragraph', 'Block Text'):
    if style_name in [s.name for s in d.styles]:
        d.styles[style_name].font.size = BODY_SIZE

for i, p in enumerate(d.paragraphs):
    if i < 5:
        continue  # title block already sized explicitly
    if p.style.name in HEADING_STYLE_NAMES:
        continue  # headings sized explicitly below
    for r in p.runs:
        if r.font.size is None:
            r.font.size = BODY_SIZE

for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = TABLE_SIZE

# ---- Master WBS & Gantt table: denser font, tighter cell margins (97 rows, 9 narrow columns) ----
# Landscape orientation (see section-split below) gives this table ~34% more
# width than it had in portrait, so the font can go back up close to TABLE_SIZE.
WBS_TABLE_SIZE = Pt(10.5)
for t in d.tables:
    if not is_wbs_gantt_table(t):
        continue
    for row in t.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side, val in (('top', 20), ('bottom', 20), ('left', 40), ('right', 40)):
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = WBS_TABLE_SIZE

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — The Complete User Guide'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

# ---- Landscape section for Part 1B.2-1B.3 (the calendar, the eight phase
# ---- tables, the Sprint/Software-Build indexes, and the Master WBS & Gantt
# ---- table) -- by far the widest, densest tables in the guide. A sectPr
# ---- embedded in a paragraph's pPr closes the section ending at that
# ---- paragraph; the very next paragraph starts the next section under
# ---- whichever sectPr closes it (or the document's trailing body sectPr,
# ---- for the final section). So: one landscape-closing copy is planted on
# ---- the last paragraph before "1B.4" (ending the landscape section), and
# ---- one portrait-closing copy is planted on the last paragraph before
# ---- "1B.2" (closing the preceding portrait section with unchanged
# ---- properties) -- everything after "1B.4" then falls through to the
# ---- document's original trailing sectPr, still portrait.
def find_heading_para_index(body_children, prefix):
    for i, el in enumerate(body_children):
        if el.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))
            if text.strip().startswith(prefix):
                return i
    return None


def nearest_preceding_paragraph(body_children, idx):
    """Walks back from body_children[idx] (exclusive) to the nearest real <w:p> --
    pandoc's auto-generated heading bookmarks (bookmarkStart/bookmarkEnd) sit
    between headings and the content before them, so a fixed offset isn't safe."""
    for j in range(idx - 1, -1, -1):
        if body_children[j].tag == qn('w:p'):
            return docx.text.paragraph.Paragraph(body_children[j], d)
    return None


def make_sectPr_copy(base_sectPr, landscape):
    new_sectPr = copy.deepcopy(base_sectPr)
    pgSz = new_sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        new_sectPr.insert(0, pgSz)
    w = int(pgSz.get(qn('w:w')) or Inches(8.5))
    h = int(pgSz.get(qn('w:h')) or Inches(11))
    if landscape:
        pgSz.set(qn('w:w'), str(max(w, h)))
        pgSz.set(qn('w:h'), str(min(w, h)))
        pgSz.set(qn('w:orient'), 'landscape')
    else:
        pgSz.set(qn('w:w'), str(min(w, h)))
        pgSz.set(qn('w:h'), str(max(w, h)))
        if pgSz.get(qn('w:orient')):
            del pgSz.attrib[qn('w:orient')]
    return new_sectPr


base_sectPr = d.sections[-1]._sectPr  # already carries the header/footer references set above
body_children = list(d.element.body)
idx_1b2 = find_heading_para_index(body_children, '1B.2 The 64-Week')
idx_1b4 = find_heading_para_index(body_children, '1B.4 Six Exception Scenarios')
p_before_1b2 = nearest_preceding_paragraph(body_children, idx_1b2) if idx_1b2 is not None else None
p_before_1b4 = nearest_preceding_paragraph(body_children, idx_1b4) if idx_1b4 is not None else None

if p_before_1b2 is not None and p_before_1b4 is not None:
    portrait_close = make_sectPr_copy(base_sectPr, landscape=False)
    landscape_close = make_sectPr_copy(base_sectPr, landscape=True)
    p_before_1b2._p.get_or_add_pPr().append(portrait_close)
    p_before_1b4._p.get_or_add_pPr().append(landscape_close)
    print('landscape section inserted: 1B.2 through end of 1B.3')
else:
    print('WARNING: could not locate landscape section boundaries -- skipping (found start:', p_before_1b2 is not None, ', end:', p_before_1b4 is not None, ')')

# ---- Force Word to recompute TOC / PAGE fields automatically on open ----
# (native TOC field ships with no cached page numbers; without this the user
# would need to manually right-click -> Update Field or press F9)
settings = d.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

d.save('mg-styled.docx')
print('wrote mg-styled.docx')