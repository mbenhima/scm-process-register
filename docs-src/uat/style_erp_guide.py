# -*- coding: utf-8 -*-
import re
import warnings
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
H4_COLOR = RGBColor(0x5A, 0x8F, 0x89)
HEADER_FILL = '1F4B45'

d = docx.Document('erp_guide-pandoc.docx')

# ---- Title block (paragraphs 0-8) ----
title_sizes = {0: 15, 1: 26, 2: 16, 3: 11.5, 4: 11.5, 5: 11.5, 6: 11, 7: 11, 8: 10.5}
title_colors = {1: H1_COLOR}
title_italics = {3, 4, 5}
for i, p in enumerate(d.paragraphs[0:9]):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(title_sizes.get(i, 11))
        r.font.color.rgb = title_colors.get(i, RGBColor(0x33, 0x33, 0x33))
        r.bold = i in (0, 1, 2)
        r.italic = i in title_italics
d.paragraphs[1].space_before = Pt(24)
d.paragraphs[8].space_after = Pt(24)

# ---- Heading-like bold paragraphs, detected by numbering pattern ----
H1_RE = re.compile(r'^(\d{1,2})\.\s')
H1_NAMED_RE = re.compile(r'^(Executive Summary|Appendix\b.*|Table of Contents)$')
H2_RE = re.compile(r'^(\d{1,2})\.(\d{1,2})\s')
TASK_RE = re.compile(r'^Task\s+\d+\s*[-—]')
STEP_RE = re.compile(r'^Step\s+\d+\s*[-—]')

for p in d.paragraphs:
    text = p.text.strip()
    if not text or not p.runs:
        continue
    is_bold_para = all((r.bold or not r.text.strip()) for r in p.runs)
    if not is_bold_para:
        continue
    level = None
    if H1_RE.match(text) or H1_NAMED_RE.match(text):
        level = 1
    elif H2_RE.match(text):
        level = 2
    elif TASK_RE.match(text):
        level = 3
    elif STEP_RE.match(text):
        level = 4
    if level is None:
        continue
    color = {1: H1_COLOR, 2: H2_COLOR, 3: H3_COLOR, 4: H4_COLOR}[level]
    size = {1: 17, 2: 13.5, 3: 12, 4: 11}[level]
    p.space_before = Pt({1: 18, 2: 14, 3: 10, 4: 8}[level])
    p.space_after = Pt(6)
    for r in p.runs:
        r.bold = True
        r.font.color.rgb = color
        r.font.size = Pt(size)

# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — Running an ERP Implementation on journi (POWERACT Consulting)'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

d.save('erp_guide-styled.docx')
print('wrote erp_guide-styled.docx')
