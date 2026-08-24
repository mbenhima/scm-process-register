# -*- coding: utf-8 -*-
import warnings
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
HEADER_FILL = '1F4B45'

d = docx.Document('uat2-pandoc.docx')

# ---- Title block (paragraphs 0-2: H1 title, subtitle, prepared-for line) ----
titles = d.paragraphs[0:3]
sizes = [32, 15, 12]
bolds = [True, False, False]
italics = [False, True, False]
for p, size, bold, italic in zip(titles, sizes, bolds, italics):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = H1_COLOR if size == 32 else RGBColor(0x33, 0x33, 0x33)
d.paragraphs[0].space_before = Pt(120)
d.paragraphs[0].style = d.styles['Title'] if 'Title' in [s.name for s in d.styles] else d.paragraphs[0].style

# ---- Heading colors ----
for p in d.paragraphs:
    style = p.style.name
    color = {'Heading 1': H1_COLOR, 'Heading 2': H2_COLOR, 'Heading 3': H3_COLOR}.get(style)
    if color:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color

# ---- Table header-row shading + borders ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — UAT Test Plan & Test Cases v2.0'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

d.save('uat2-styled.docx')
print('wrote uat2-styled.docx')
