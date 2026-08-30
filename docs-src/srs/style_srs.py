# -*- coding: utf-8 -*-
import warnings
import docx
from docx.shared import RGBColor, Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
H4_COLOR = RGBColor(0x5A, 0x8F, 0x89)
HEADER_FILL = '1F4B45'

d = docx.Document('srs-pandoc.docx')

# ---- The TOCPLACEHOLDERXYZ paragraph (index 5) is where update_toc.py will
# ---- insert a genuine LibreOffice ContentIndex later in the pipeline; wrap
# ---- it in page breaks here so it lands on its own page regardless of how
# ---- long the eventual index renders.
def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p

placeholder_p = None
for p in d.paragraphs:
    if p.text.strip() == 'TOCPLACEHOLDERXYZ':
        placeholder_p = p
        break
if placeholder_p is not None:
    pb_before = make_page_break_para()
    pb_after = make_page_break_para()
    placeholder_p._p.addprevious(pb_before)
    placeholder_p._p.addnext(pb_after)

# ---- Margins ----
for section in d.sections:
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.page_height is None:
        section.page_height = Inches(11)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
USABLE_WIDTH = d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin

# ---- Title block (paragraphs 0-4) ----
titles = d.paragraphs[0:5]
sizes = [30, 15, 11.5, 11.5, 11]
bolds = [True, False, False, False, False]
italics = [False, False, True, True, False]
for p, size, bold, italic in zip(titles, sizes, bolds, italics):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = H1_COLOR if size == 30 else RGBColor(0x33, 0x33, 0x33)
d.paragraphs[0].space_before = Pt(100)
d.paragraphs[4].space_after = Pt(20)

# ---- Heading colors and sizes ----
HEADING_SPECS = {
    'Heading 1': (H1_COLOR, Pt(21)),
    'Heading 2': (H1_COLOR, Pt(19)),
    'Heading 3': (H2_COLOR, Pt(15)),
    'Heading 4': (H3_COLOR, Pt(13)),
}
for p in d.paragraphs:
    spec = HEADING_SPECS.get(p.style.name)
    if spec:
        color, size = spec
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color
            r.font.size = size
        if p.style.name in ('Heading 3', 'Heading 4'):
            p.space_before = Pt(14)

# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

# ---- Body/table font sizing ----
BODY_SIZE = Pt(11.5)
TABLE_SIZE = Pt(9.5)
HEADING_STYLE_NAMES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title'}

for style_name in ('Normal', 'Body Text', 'Compact', 'First Paragraph', 'List Paragraph', 'Block Text', 'Source Code'):
    if style_name in [s.name for s in d.styles]:
        d.styles[style_name].font.size = BODY_SIZE

for i, p in enumerate(d.paragraphs):
    if i < 5:
        continue
    if p.style.name in HEADING_STYLE_NAMES:
        continue
    for r in p.runs:
        if r.font.size is None:
            r.font.size = BODY_SIZE

for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = TABLE_SIZE

# ---- Prevent a table row's content from splitting across a page break ----
for t in d.tables:
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — Software Requirements Specification'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

# ---- Force Word to recompute TOC / PAGE fields automatically on open ----
settings = d.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

d.save('srs-styled.docx')
print('wrote srs-styled.docx')
