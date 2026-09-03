# -*- coding: utf-8 -*-
import copy
import warnings
import docx
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
H4_COLOR = RGBColor(0x5A, 0x8F, 0x89)
HEADER_FILL = '1F4B45'

d = docx.Document('cc_mfg-pandoc.docx')


# ---- Move the native TOC field from the very top of the document (where
# ---- pandoc puts it) to after the title block.
def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


body = d.element.body
sdt = body.find(qn('w:sdt'))
if sdt is not None:
    body.remove(sdt)
    title_end_p = d.paragraphs[5]._p  # 'Version 1.0 ...' -- last title-block paragraph
    pb_before_toc = make_page_break_para()
    pb_after_toc = make_page_break_para()
    title_end_p.addnext(pb_before_toc)
    pb_before_toc.addnext(sdt)
    sdt.addnext(pb_after_toc)

# ---- Margins ----
for section in d.sections:
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.page_height is None:
        section.page_height = Inches(11)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
USABLE_WIDTH = d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin
LANDSCAPE_USABLE_WIDTH = d.sections[0].page_height - d.sections[0].left_margin - d.sections[0].right_margin

# ---- Title block (paragraphs 0-5) ----
title_sizes = {0: 13, 1: 24, 2: 13, 3: 11, 4: 10.5, 5: 10.5}
title_bolds = {0: True, 1: True}
for i, p in enumerate(d.paragraphs[0:6]):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(title_sizes.get(i, 11))
        r.font.color.rgb = H1_COLOR if i == 1 else RGBColor(0x33, 0x33, 0x33)
        r.bold = title_bolds.get(i, False)
d.paragraphs[0].space_before = Pt(90)
d.paragraphs[1].space_before = Pt(18)
d.paragraphs[5].space_after = Pt(24)

# ---- Heading colors and sizes ----
HEADING_SPECS = {
    'Heading 2': (H1_COLOR, Pt(19)),
    'Heading 3': (H2_COLOR, Pt(14.5)),
    'Heading 4': (H3_COLOR, Pt(12.5)),
}
for p in d.paragraphs:
    spec = HEADING_SPECS.get(p.style.name)
    if spec:
        color, size = spec
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color
            r.font.size = size
        if p.style.name == 'Heading 2':
            p.space_before = Pt(20)
        elif p.style.name == 'Heading 3':
            p.space_before = Pt(14)
        else:
            p.space_before = Pt(10)

# ---- Bold "Detailed description / Trigger / ..." labels a touch larger ----
for p in d.paragraphs:
    text = p.text.strip()
    if not text or not p.runs:
        continue
    is_bold_lead = p.runs[0].bold and text.split('.')[0].endswith('*') is False and any(
        text.startswith(prefix) for prefix in (
            'Detailed description', 'Trigger', 'Timeline impact', 'Recovery tasks', 'Outputs', 'RACSI'))
    if is_bold_lead:
        for r in p.runs:
            r.font.color.rgb = H3_COLOR


# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Prevent a table row's content from splitting across a page break ----
for t in d.tables:
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

# ---- Base body / table font sizing ----
BODY_SIZE = Pt(11)
TABLE_SIZE = Pt(9.5)
HEADING_STYLE_NAMES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'}

for style_name in ('Normal', 'Body Text', 'Compact', 'First Paragraph', 'List Paragraph', 'Block Text'):
    if style_name in [s.name for s in d.styles]:
        d.styles[style_name].font.size = BODY_SIZE

for i, p in enumerate(d.paragraphs):
    if i < 6:
        continue  # title block already sized explicitly
    if p.style.name in HEADING_STYLE_NAMES:
        continue
    for r in p.runs:
        if r.font.size is None:
            r.font.size = BODY_SIZE

for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = TABLE_SIZE

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — Leading a Cultural & Values Transformation (One Bouregreg)'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

# ---- Landscape section for Part 4 (all of it) through end of document ----
# Part 4 is the last Part in this guide and carries every wide table (the
# weekly PM/CM tables and both Master WBS & Gantt tables), so a single
# section break before "Part 4" that runs landscape to the end of the
# document is sufficient -- no need to switch back to portrait afterward.
def find_heading_para_index(body_children, prefix):
    for i, el in enumerate(body_children):
        if el.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))
            if text.strip().startswith(prefix):
                return i
    return None


def nearest_preceding_paragraph(body_children, idx):
    for j in range(idx - 1, -1, -1):
        if body_children[j].tag == qn('w:p'):
            return docx.text.paragraph.Paragraph(body_children[j], d)
    return None


def make_sectPr_copy(base_sectPr, landscape):
    new_sectPr = copy.deepcopy(base_sectPr)
    pgSz = new_sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        new_sectPr.insert(0, pgSz)
    w = int(pgSz.get(qn('w:w')) or Inches(8.5))
    h = int(pgSz.get(qn('w:h')) or Inches(11))
    if landscape:
        pgSz.set(qn('w:w'), str(max(w, h)))
        pgSz.set(qn('w:h'), str(min(w, h)))
        pgSz.set(qn('w:orient'), 'landscape')
    else:
        pgSz.set(qn('w:w'), str(min(w, h)))
        pgSz.set(qn('w:h'), str(max(w, h)))
        if pgSz.get(qn('w:orient')):
            del pgSz.attrib[qn('w:orient')]
    return new_sectPr


base_sectPr = d.sections[-1]._sectPr  # the document's TRAILING sectPr -- governs
# from the last section break to the end of the document. A sectPr embedded in
# a paragraph's pPr closes the section ENDING at that paragraph; whatever comes
# after falls under the next break, or this trailing sectPr if none follows.
body_children = list(d.element.body)
idx_part4 = find_heading_para_index(body_children, 'Part 4')
p_before_part4 = nearest_preceding_paragraph(body_children, idx_part4) if idx_part4 is not None else None

if p_before_part4 is not None:
    # Close the section before Part 4 as portrait (a copy, unchanged orientation).
    portrait_close = make_sectPr_copy(base_sectPr, landscape=False)
    p_before_part4._p.get_or_add_pPr().append(portrait_close)
    # Flip the document's own trailing sectPr to landscape IN PLACE -- this is
    # the section that now runs from Part 4 to the end of the document.
    pgSz = base_sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        base_sectPr.insert(0, pgSz)
    w = int(pgSz.get(qn('w:w')) or Inches(8.5))
    h = int(pgSz.get(qn('w:h')) or Inches(11))
    pgSz.set(qn('w:w'), str(max(w, h)))
    pgSz.set(qn('w:h'), str(min(w, h)))
    pgSz.set(qn('w:orient'), 'landscape')
    print('landscape section inserted: Part 4 through end of document')
else:
    print('WARNING: could not locate Part 4 heading -- skipping landscape section')

# ---- Force Word to recompute TOC / PAGE fields automatically on open ----
settings = d.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

d.save('cc_mfg-styled.docx')
print('wrote cc_mfg-styled.docx')
