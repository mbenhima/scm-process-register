# -*- coding: utf-8 -*-
import warnings
import docx
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore', category=UserWarning)

H1_COLOR = RGBColor(0x1F, 0x4B, 0x45)
H2_COLOR = RGBColor(0x27, 0x56, 0x50)
H3_COLOR = RGBColor(0x3F, 0x82, 0x7B)
HEADER_FILL = '1F4B45'

d = docx.Document('cc_guide-pandoc.docx')


# ---- Move the native TOC field from the very top of the document (where
# ---- pandoc puts it) to after the title block, so reading order is
# ---- Cover -> TOC -> Executive Summary, not TOC -> Cover. Page breaks pin
# ---- the cover to its own page and the TOC to its own page range.
def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


body = d.element.body
sdt = body.find(qn('w:sdt'))
if sdt is not None:
    body.remove(sdt)
    title_end_p = d.paragraphs[7]._p  # 'Version 1.0 ...' -- last title-block paragraph
    pb_before_toc = make_page_break_para()
    pb_after_toc = make_page_break_para()
    title_end_p.addnext(pb_before_toc)
    pb_before_toc.addnext(sdt)
    sdt.addnext(pb_after_toc)

# ---- Margins ----
for section in d.sections:
    if section.page_width is None:
        section.page_width = Inches(8.5)
    if section.page_height is None:
        section.page_height = Inches(11)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

# ---- Title block (paragraphs 0-7) ----
title_sizes = {0: 13, 1: 24, 2: 15, 3: 11, 4: 11, 5: 11, 6: 10.5, 7: 10.5}
title_bolds = {0: True, 1: True}
title_italics = {3, 4}
for i, p in enumerate(d.paragraphs[0:8]):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(title_sizes.get(i, 11))
        r.font.color.rgb = H1_COLOR if i == 1 else RGBColor(0x33, 0x33, 0x33)
        r.bold = title_bolds.get(i, False)
        r.italic = i in title_italics
d.paragraphs[0].space_before = Pt(90)
d.paragraphs[1].space_before = Pt(18)
d.paragraphs[7].space_after = Pt(24)

# ---- Heading colors and sizes ----
HEADING_SPECS = {
    'Heading 2': (H1_COLOR, Pt(19)),
    'Heading 3': (H2_COLOR, Pt(14)),
}
for p in d.paragraphs:
    spec = HEADING_SPECS.get(p.style.name)
    if spec:
        color, size = spec
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = color
            r.font.size = size
        if p.style.name == 'Heading 2':
            p.space_before = Pt(20)
        else:
            p.space_before = Pt(12)

# ---- Bold "Task"/"Step"/"Exception"/"SIPOC"-caption paragraphs a touch larger ----
for p in d.paragraphs:
    text = p.text.strip()
    if not text or not p.runs:
        continue
    is_bold_para = all((r.bold or not r.text.strip()) for r in p.runs)
    if not is_bold_para:
        continue
    if text.startswith('Task ') or text.startswith('Step '):
        for r in p.runs:
            r.font.color.rgb = H3_COLOR
            r.font.size = Pt(12)
        p.space_before = Pt(10)
        p.space_after = Pt(4)
    elif text in ('SIPOC', "Tasks, Steps, Techniques & RACSI"):
        for r in p.runs:
            r.font.color.rgb = H3_COLOR
            r.font.size = Pt(11.5)
        p.space_before = Pt(8)


# ---- Table header-row shading ----
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


for t in d.tables:
    if len(t.rows) == 0:
        continue
    header = t.rows[0]
    for cell in header.cells:
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- Prevent a table row's content from splitting across a page break ----
for t in d.tables:
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

# ---- Base body / table font sizing ----
BODY_SIZE = Pt(11)
TABLE_SIZE = Pt(10)
HEADING_STYLE_NAMES = {'Heading 1', 'Heading 2', 'Heading 3'}

for style_name in ('Normal', 'Body Text', 'Compact', 'First Paragraph', 'List Paragraph', 'Block Text'):
    if style_name in [s.name for s in d.styles]:
        d.styles[style_name].font.size = BODY_SIZE

for i, p in enumerate(d.paragraphs):
    if i < 8:
        continue  # title block already sized explicitly
    if p.style.name in HEADING_STYLE_NAMES:
        continue
    for r in p.runs:
        if r.font.size is None:
            r.font.size = BODY_SIZE

for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.size is None or r.font.size == BODY_SIZE:
                        r.font.size = TABLE_SIZE

# ---- Header / footer with page numbers ----
section = d.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = 'journi — Leading a Cultural & Values Transformation (POWERACT Consulting)'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run('Page ')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_sep = OxmlElement('w:fldChar')
fld_sep.set(qn('w:fldCharType'), 'separate')
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
r2 = fp.add_run()
r2.font.size = Pt(8)
r2._r.append(fld_begin)
r2._r.append(instr)
r2._r.append(fld_sep)
r2._r.append(fld_end)

# ---- Force Word to recompute TOC / PAGE fields automatically on open ----
settings = d.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

d.save('cc_guide-styled.docx')
print('wrote cc_guide-styled.docx')
